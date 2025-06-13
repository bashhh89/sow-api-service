# main.py
# FINAL VERSION
# This API receives a workspace and chat ID, fetches the last message from AnythingLLM,
# converts it to a DOCX using a template, uploads it, and returns a download link.

import io
import os
import re
import requests
from fastapi import FastAPI, Response, HTTPException
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt
from docx.table import _Cell

# --- CONFIGURATION ---
# These will be loaded from environment variables in your Easypanel setup
ANYTHINGLLM_API_URL = os.environ.get("ANYTHINGLLM_API_URL") # e.g., "http://192.168.1.5:3001"
ANYTHINGLLM_API_KEY = os.environ.get("ANYTHINGLLM_API_KEY")

# Initialize the FastAPI app
app = FastAPI(
    title="Smart SOW Document Converter",
    description="An API that pulls chat history from AnythingLLM to generate a .docx file.",
    version="2.0.0",
)

# --- Pydantic Models ---
class ConversionRequest(BaseModel):
    workspace_slug: str
    chat_id: str
    filename: str = "SOW-Document.docx"

class ConversionResponse(BaseModel):
    status: str
    download_url: str

# --- Helper Functions (No changes here) ---
def parse_and_add_paragraph(paragraph_text: str, document: Document):
    p = document.add_paragraph()
    pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)'
    parts = re.split(pattern, paragraph_text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = p.add_run(part[3:-3]); run.bold = True; run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2]); run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1]); run.italic = True
        elif part: p.add_run(part)

def create_docx_from_markdown(markdown_text: str, document: Document):
    lines = markdown_text.split('\n')
    in_table, table_data, i = False, [], 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            if not in_table: in_table = True
            table_data.append([cell.strip() for cell in line.strip('|').split('|')])
            i += 1; continue
        elif in_table:
            in_table = False
            if table_data:
                headers = table_data[0]
                records = table_data[2:] if len(table_data) > 1 and '---' in table_data[1][0] else table_data[1:]
                table = document.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    hdr_cells[j].text = header; hdr_cells[j].paragraphs[0].runs[0].bold = True
                for record in records:
                    row_cells = table.add_row().cells
                    for j, cell_text in enumerate(record): row_cells[j].text = cell_text
            table_data = []
        if i < len(lines) and lines[i].strip().startswith('|'): continue
        if line.startswith('# '): document.add_heading(line[2:], level=1)
        elif line.startswith('## '): document.add_heading(line[3:], level=2)
        elif line.startswith('### '): document.add_heading(line[4:], level=3)
        elif line.startswith('* '):
            parse_and_add_paragraph(line[2:], document); document.paragraphs[-1].style = 'List Bullet'
        elif line.strip() == '---': document.add_page_break()
        elif line.strip(): parse_and_add_paragraph(line, document)
        i += 1

# --- The Main API Endpoint ---
@app.post("/generate-from-chat", response_model=ConversionResponse)
async def generate_from_chat(request: ConversionRequest):
    if not ANYTHINGLLM_API_URL or not ANYTHINGLLM_API_KEY:
        raise HTTPException(status_code=500, detail="Server is not configured with AnythingLLM API credentials.")

    # Step 1: Fetch chat history from AnythingLLM API
    history_url = f"{ANYTHINGLLM_API_URL}/api/v1/workspace/{request.workspace_slug}/chat/{request.chat_id}/history"
    headers = {"Authorization": f"Bearer {ANYTHINGLLM_API_KEY}"}
    try:
        response = requests.get(history_url, headers=headers, timeout=15)
        response.raise_for_status()
        history = response.json()
    except requests.exceptions.RequestException as e:
        raise HTTPException(status_code=502, detail=f"Failed to fetch chat history from AnythingLLM: {e}")

    # Find the last message from the 'assistant' that isn't an error message
    markdown_text = ""
    for message in reversed(history.get('history', [])):
        if message.get('role') == 'assistant' and "unable to directly export" not in message.get('text', '').lower():
            markdown_text = message.get('text')
            break
            
    if not markdown_text:
        raise HTTPException(status_code=404, detail="Could not find a suitable SOW message in the chat history.")

    # Step 2: Create the DOCX file
    template_path = 'template.docx'
    document = Document(template_path) if os.path.exists(template_path) else Document()
    create_docx_from_markdown(markdown_text, document)

    # Step 3: Upload the file
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    try:
        server_response = requests.get('https://api.gofile.io/getServer', timeout=10)
        server = server_response.json()['data']['server']
        upload_url = f'https://{server}.gofile.io/uploadFile'
        files = {'file': (request.filename, file_stream.read(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        upload_response = requests.post(upload_url, files=files, timeout=30)
        upload_data = upload_response.json()
        if upload_data.get("status") == "ok":
            return {"status": "success", "download_url": upload_data['data']['downloadPage']}
        else:
            raise HTTPException(status_code=500, detail="File hosting service failed.")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"File upload error: {e}")

